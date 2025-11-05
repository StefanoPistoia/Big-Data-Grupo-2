# -*- coding: utf-8 -*-
"""
Created on Tue Oct 21 09:35:07 2025 (actualizado)

@author: agostina.giovanardi
"""

import pandas as pd
from sklearn.model_selection import train_test_split
from scipy import stats
from docx import Document
import statsmodels.api as sm
import numpy as np
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from sklearn.neighbors import KNeighborsClassifier
from sklearn.preprocessing import StandardScaler
from sklearn.metrics import accuracy_score, classification_report

# Cargar el archivo Excel en un DataFrame de pandas
df = pd.read_excel("db_nea_respuestas.xlsx")

# Cuando ESTADO es 2 (desocupado), 3 (inactivo) o 4 (menor), las horas trabajadas son 0.
# Se imputa el valor 0 en 'horastrab' para estos casos.
df.loc[df['ESTADO'].isin([2, 3, 4]), 'horastrab'] = 0

# Nombre de la columna que contiene el año
columna_anio = 'ANO4'
# Semilla para la reproducibilidad
random_seed = 444

# --- Procesamiento para el año 2005 ---
respondieron_2005 = df[df[columna_anio] == 2005]

train_2005, test_2005 = train_test_split(
    respondieron_2005,
    test_size=0.3,
    random_state=random_seed
)

# --- Procesamiento para el año 2025 ---

respondieron_2025 = df[df[columna_anio] == 2025]

train_2025, test_2025 = train_test_split(
    respondieron_2025,
    test_size=0.3,
    random_state=random_seed
)

# Definimos variables
Y = 'pobre'
X = ['horastrab','educ','edad','edad2','cobertura_medica','sexo','ESTADO','estado_civil']

# --- Tabla de diferencia de medias con p-values entre train y test para cada año ---

results = []

# Diccionario para mapear el año a sus dataframes de train y test
dataframes_por_anio = {
    2005: {'train': train_2005, 'test': test_2005},
    2025: {'train': train_2025, 'test': test_2025}
}

print("\nCalculando diferencias de medias y p-values...")

for anio, dfs in dataframes_por_anio.items():
    train_df = dfs['train']
    test_df = dfs['test']
    
    print(f"  Procesando año: {anio}")
    
    for var in X:
        # Asegurarse de que la variable existe en ambos dataframes
        if var in train_df.columns and var in test_df.columns:
            # Calcular medias
            mean_train = train_df[var].mean()
            mean_test = test_df[var].mean()
            
            # Calcular diferencia de medias
            diff_mean = mean_train - mean_test
            
            # Realizar t-test de Student para muestras independientes
            # Usamos .dropna() para manejar posibles valores faltantes
            # equal_var=False para Welch's t-test, que no asume varianzas iguales
            t_stat, p_value = stats.ttest_ind(
                train_df[var].dropna(), 
                test_df[var].dropna(), 
                equal_var=False
            )
            
            results.append({
                'Año': anio,
                'Variable': var,
                'Media_Train': mean_train,
                'Media_Test': mean_test,
                'Diferencia_Medias': diff_mean,
                'P_Value': p_value
            })
        else:
            print(f"    Advertencia: La variable '{var}' no se encontró en los dataframes del año {anio}. Se omitirá.")

# Crear un DataFrame con los resultados
df_medias_pvalues = pd.DataFrame(results)

# Mostrar el DataFrame resultante
print("\nTabla de Diferencia de Medias y P-values:")
print(df_medias_pvalues)

# Exportar a Excel
df_medias_pvalues.to_excel("diferencia_medias_train_test.xlsx", index=False)





#fijarse que categorías son base para las dummies





# --- Regresión logística ---
train_2025_full = train_2025[[Y] + X].dropna()

# 1. Preparar los datos de entrenamiento para el modelo
y_train = train_2025_full[Y]
X_train = train_2025_full[X]

# --- Diagnóstico: Verificar valores únicos en las variables categóricas ---
print("\nVerificando valores únicos en X_train antes de crear dummies:")
print(X_train[categorical_vars].nunique())

# Convertir variables categóricas en dummies
categorical_vars = ['cobertura_medica', 'sexo', 'ESTADO', 'estado_civil']
X_train_dummies = pd.get_dummies(X_train, columns=categorical_vars, drop_first=True, dtype=float)

# Statsmodels requiere que se añada una constante manualmente
X_train_const = sm.add_constant(X_train_dummies)

# 2. Instanciar y entrenar el modelo Logit
logit_model = sm.Logit(y_train, X_train_const)
result = logit_model.fit()

# 3. Resumen del modelo
print(result.summary())

# Extraer resultados
coef = result.params
se = result.bse
z = result.tvalues
p = result.pvalues
conf = result.conf_int()
conf.columns = ['IC 2.5%', 'IC 97.5%']

# Crear DataFrame
tabla_resultados = pd.DataFrame({
    'Variable': coef.index,
    'Coef. (β)': coef.values,
    'exp(β)': np.exp(coef.values),
    'Error Std.': se.values,
    'z': z.values,
    'p-value': p.values,
    'IC 2.5%': conf['IC 2.5%'].values,
    'IC 97.5%': conf['IC 97.5%'].values
})

tabla_resultados = tabla_resultados.round(4)

# --- Exportar a Word ---
doc = Document()
doc.add_heading('Resultados del modelo Logit', level=1)

# Crear tabla
t = doc.add_table(rows=1, cols=len(tabla_resultados.columns))
t.style = 'Light List Accent 1'

# Cabeceras
hdr_cells = t.rows[0].cells
for i, col in enumerate(tabla_resultados.columns):
    hdr_cells[i].text = col

# Filas (con negrita si p < 0.05)
for _, row in tabla_resultados.iterrows():
    row_cells = t.add_row().cells
    for j, val in enumerate(row):
        cell = row_cells[j].paragraphs[0].add_run(str(val))
        if j == 0 or tabla_resultados.loc[_, 'p-value'] < 0.05:
            # Si es nombre de variable o p < 0.05 => negrita
            cell.bold = True

# Ajustar tamaño de fuente
for row in t.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

# Guardar documento
doc.save('Resultados_Logit.docx')

import matplotlib.pyplot as plt


# 1. Crear rango de valores para 'educ'
educ_range = np.linspace(X_train_const['educ'].min(), X_train_const['educ'].max(), 100)

# 2. Crear un DataFrame con 'educ' variable y los demás fijos en su media
X_pred = X_train_const.copy()
for col in X_pred.columns:
    if col != 'educ':
        X_pred[col] = X_pred[col].mean()

# Reemplazamos 'educ' por el rango deseado
X_pred = X_pred.loc[X_pred.index.repeat(len(educ_range))].reset_index(drop=True)
X_pred['educ'] = np.tile(educ_range, len(X_train_const))

# 3. Predecir probabilidades
pred_probs = result.predict(X_pred)

# 4. Graficar relación entre educación y probabilidad predicha
plt.figure(figsize=(8, 5))
plt.plot(educ_range, pred_probs.groupby(X_pred['educ']).mean(), color='darkblue', linewidth=2)
plt.title('Probabilidad predicha de ser pobre según nivel educativo', fontsize=13)
plt.xlabel('Años de educación', fontsize=12)
plt.ylabel('Probabilidad predicha de ser pobre', fontsize=12)
plt.grid(alpha=0.3)
plt.ylim(0, 1)
plt.show()
















# --- Clasificación con K-Nearest Neighbors (KNN) ---

# Usamos las mismas variables X e Y que en la regresión logística
train_df_knn = train_2025[[Y] + X].dropna()
test_df_knn = test_2025[[Y] + X].dropna()

y_train_knn = train_df_knn[Y]
X_train_knn = train_df_knn[X]

y_test_knn = test_df_knn[Y]
X_test_knn = test_df_knn[X]

# Convertir variables categóricas en dummies
X_train_knn_dummies = pd.get_dummies(X_train_knn, columns=categorical_vars, drop_first=True, dtype=float)
X_test_knn_dummies = pd.get_dummies(X_test_knn, columns=categorical_vars, drop_first=True, dtype=float)

# Alinear columnas para asegurar que train y test tengan las mismas 
X_train_aligned, X_test_aligned = X_train_knn_dummies.align(X_test_knn_dummies, join='inner', axis=1, fill_value=0)

# 2. Escalar las características
# KNN es sensible a la escala de los datos, por lo que estandarizamos las variables.
scaler = StandardScaler()
X_train_scaled = scaler.fit_transform(X_train_aligned)
X_test_scaled = scaler.transform(X_test_aligned)

# 3. Iterar sobre los valores de K, entrenar y evaluar el modelo
k_values = [1, 5, 10]

for k in k_values:
    print(f"\n--- Resultados para K = {k} ---")
    
    # Instanciar el modelo
    knn = KNeighborsClassifier(n_neighbors=k)
    
    # Entrenar el modelo
    knn.fit(X_train_scaled, y_train_knn)
    
    # Predecir en el conjunto de prueba
    y_pred_knn = knn.predict(X_test_scaled)
    
    # Evaluar el modelo
    accuracy = accuracy_score(y_test_knn, y_pred_knn)
    print(f"Accuracy: {accuracy:.4f}")
    print("Reporte de Clasificación:")
    print(classification_report(y_test_knn, y_pred_knn, target_names=['No Pobre (0)', 'Pobre (1)']))
    
# --- Búsqueda del K Óptimo con Cross-Validation (KNN con K-CV) ---

from sklearn.model_selection import cross_val_score

print("\n--- Buscando el K óptimo para KNN con Cross-Validation (5-fold) ---")

# 1. Definir el rango de valores de K para probar
k_range = range(1, 11)
cv_scores = []

# 2. Realizar validación cruzada de 5 folds para cada valor de K
print("Calculando el accuracy promedio para cada K...")
for k in k_range:
    knn_cv = KNeighborsClassifier(n_neighbors=k)
    # cross_val_score divide los datos, entrena y evalúa automáticamente
    # Usamos X_train_scaled y y_train_knn, que son nuestros datos de entrenamiento
    scores = cross_val_score(knn_cv, X_train_scaled, y_train_knn, cv=5, scoring='accuracy')
    cv_scores.append(scores.mean())

# 3. Encontrar el K óptimo
optimal_k = k_range[np.argmax(cv_scores)]
max_accuracy = max(cv_scores)


# 4. Graficar los resultados
plt.figure(figsize=(10, 6))
plt.plot(k_range, cv_scores, marker='o', linestyle='--', color='b', label='Accuracy de CV')
plt.axvline(optimal_k, color='r', linestyle='-', label=f'K Óptimo = {optimal_k}')
plt.title('Accuracy de KNN vs. Número de Vecinos (K) - Cross-Validation', fontsize=14)
plt.xlabel('Número de Vecinos (K)', fontsize=12)
plt.ylabel('Accuracy Promedio (5-Fold CV)', fontsize=12)
plt.xticks(k_range)
plt.legend()
plt.grid(True)
plt.show()

# 5. Comentario sobre el resultado
print("\n--- Comentario sobre el K Óptimo ---")
print(f"El gráfico muestra el rendimiento del modelo para diferentes valores de K. Un K bajo (como 1) puede llevar a un sobreajuste, ya que el modelo es muy sensible a puntos individuales y ruido. A medida que K aumenta, el modelo se generaliza mejor, y el accuracy tiende a subir.")
print(f"En este caso, el accuracy más alto se alcanza con K={optimal_k}. A partir de este punto, aumentar K podría hacer que el modelo sea demasiado simple (underfitting), perdiendo detalles importantes y disminuyendo su rendimiento.")
print(f"Por lo tanto, {optimal_k} es el número óptimo de vecinos cercanos para este problema, ya que representa el mejor balance entre sesgo y varianza según la validación cruzada.")


# --- Regularización con LASSO (L1) y Ridge (L2) ---

from sklearn.linear_model import LogisticRegression

# 1. Definir la grilla de parámetros de penalidad
lambdas = np.logspace(-5, 5, 11)

# En sklearn, C es el inverso de la fuerza de penalidad (C = 1/λ)
cs = 1 / lambdas

coefs_lasso = []
coefs_ridge = []

# Usamos los datos escalados del entrenamiento de KNN, que ya están preparados

# Iteramos para LASSO
for c in cs:
    # solver='liblinear' es una buena opción para datasets pequeños y funciona con L1 y L2
    model_lasso = LogisticRegression(penalty='l1', C=c, solver='liblinear', random_state=random_seed)
    model_lasso.fit(X_train_scaled, y_train_knn)
    coefs_lasso.append(model_lasso.coef_[0])

# Iteramos para Ridge
for c in cs:
    model_ridge = LogisticRegression(penalty='l2', C=c, solver='liblinear', random_state=random_seed)
    model_ridge.fit(X_train_scaled, y_train_knn)
    coefs_ridge.append(model_ridge.coef_[0])

# Graficamos los coeficientes en dos paneles
fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(20, 8), sharey=True)

# Nombres de las variables para las leyendas
feature_names = X_train_aligned.columns

# Panel para LASSO 
ax1.plot(np.log10(lambdas), coefs_lasso)
ax1.set_xlabel('log10(λ)', fontsize=12)
ax1.set_ylabel('Coeficientes', fontsize=12)
ax1.set_title('Coeficientes de Regresión Logística con Penalidad LASSO (L1)', fontsize=14)
ax1.grid(True)
ax1.legend(feature_names, loc='upper right')

# Panel para Ridge 
ax2.plot(np.log10(lambdas), coefs_ridge)
ax2.set_xlabel('log10(λ)', fontsize=12)
ax2.set_title('Coeficientes de Regresión Logística con Penalidad Ridge (L2)', fontsize=14)
ax2.grid(True)
ax2.legend(feature_names, loc='upper right')

plt.suptitle('Evolución de Coeficientes vs. Fuerza de Penalidad (λ)', fontsize=16, y=1)
plt.show()

# 5. Interpretación de la regularización
print("\n--- Interpretación de la Regularización ---")
print("\nPenalidad LASSO (L1):")
print("El gráfico de la izquierda muestra cómo la penalidad LASSO afecta a los coeficientes del modelo. LASSO (Least Absolute Shrinkage and Selection Operator) añade un término de penalidad igual al valor absoluto de la magnitud de los coeficientes.")
print(" - Para valores bajos de λ (alta C, a la izquierda del gráfico), la penalidad es débil y los coeficientes son grandes, similar a una regresión logística sin regularización.")
print(" - A medida que λ aumenta (hacia la derecha), la penalidad se vuelve más fuerte. LASSO fuerza a que los coeficientes de las variables menos importantes se reduzcan exactamente a cero.")
print(" - Esto se puede ver claramente en el gráfico: muchas líneas de coeficientes convergen y se 'aplastan' en cero a medida que λ crece. Este comportamiento hace que LASSO sea muy útil para la selección de características (feature selection), ya que elimina efectivamente las variables menos relevantes del modelo.")

print("\nPenalidad Ridge (L2):")
print("El gráfico de la derecha muestra el efecto de la penalidad Ridge. Ridge añade un término de penalidad igual al cuadrado de la magnitud de los coeficientes.")
print(" - Al igual que con LASSO, para valores bajos de λ (a la izquierda), los coeficientes son grandes.")
print(" - A medida que λ aumenta, Ridge también reduce ('encoge') los coeficientes hacia cero para evitar el sobreajuste.")
print(" - Sin embargo, a diferencia de LASSO, la penalidad Ridge no fuerza los coeficientes a ser exactamente cero (a menos que λ sea infinito). Los coeficientes se aproximan a cero pero rara vez lo alcanzan.")
print(" - Esto es visible en el gráfico: todas las líneas de coeficientes disminuyen y tienden a cero a medida que λ aumenta, pero no se 'aplastan' en cero como en el caso de LASSO. Ridge es útil cuando se espera que todas las características contribuyan al modelo, aunque sea en pequeña medida.")

print("\nConclusión:")
print("Ambas técnicas son efectivas para combatir el sobreajuste (overfitting) al penalizar modelos complejos con coeficientes grandes. La principal diferencia práctica es que LASSO puede realizar selección de variables, mientras que Ridge retiene todas las variables, simplemente reduciendo su impacto.")


# --- Selección de λ Óptimo con Cross-Validation y Box-plots de Error ---

from sklearn.linear_model import LogisticRegressionCV

print("\n--- Selección de λ Óptimo con LogisticRegressionCV (5-fold) ---")

# Usamos la misma grilla de parámetros
# LogisticRegressionCV espera los valores de C (1/λ)

# --- Penalidad LASSO (L1) ---
print("\nProcesando para LASSO (L1)...")

# Instanciar y entrenar el modelo con CV
# El atributo 'scores_' guardará los scores de cada fold para cada valor de C
# El atributo 'coefs_' guardará los coeficientes de cada fold.
# Para acceder a los coeficientes por fold, es importante no usar refit=True.
# Por defecto, refit=False en versiones antiguas, pero es True por defecto en las nuevas.
# Para asegurar que tenemos los coeficientes de cada fold, los extraemos de un modelo
# que no se re-entrena al final.
lasso_cv_model = LogisticRegressionCV(
    Cs=cs,
    cv=5,
    penalty='l1',
    solver='liblinear',
    scoring='accuracy',
    random_state=random_seed,
    max_iter=1000, # Aumentar iteraciones para asegurar convergencia
    refit=True # Re-entrenar con el mejor C para obtener el C_ óptimo
).fit(X_train_scaled, y_train_knn)

# Para obtener los scores y coefs de cada fold, necesitamos un modelo sin refit
lasso_cv_scores_and_coefs = LogisticRegressionCV(
    Cs=cs, cv=5, penalty='l1', solver='liblinear', scoring='accuracy', random_state=random_seed, max_iter=1000, refit=False
).fit(X_train_scaled, y_train_knn)


# El atributo scores_ es un dict, tomamos los scores de la clase positiva (1)
# Su forma es (n_folds, n_cs)
lasso_errors = 1 - lasso_cv_scores_and_coefs.scores_[1]

# El atributo `coefs_paths_` es un diccionario. Para la clase positiva (1),
# contiene un array de forma (n_folds, n_cs, n_features).
# Calculamos la proporción de coeficientes cero para cada fold y cada C
prop_zeros_lasso = np.mean(lasso_cv_scores_and_coefs.coefs_paths_[1] == 0, axis=2)

lambda_optimo_lasso = 1 / lasso_cv_model.C_[0]
print(f"El λ^cv óptimo seleccionado para LASSO es: {lambda_optimo_lasso:.5f}")

# --- Penalidad Ridge (L2) ---
print("\nProcesando para Ridge (L2)...")

ridge_cv_model = LogisticRegressionCV(
    Cs=cs, cv=5, penalty='l2', solver='liblinear', scoring='accuracy', random_state=random_seed, max_iter=1000, refit=True
).fit(X_train_scaled, y_train_knn)

ridge_cv_scores = LogisticRegressionCV(
    Cs=cs, cv=5, penalty='l2', solver='liblinear', scoring='accuracy', random_state=random_seed, max_iter=1000, refit=False
).fit(X_train_scaled, y_train_knn)

ridge_errors = 1 - ridge_cv_scores.scores_[1]
lambda_optimo_ridge = 1 / ridge_cv_model.C_[0]
print(f"El λ^cv óptimo seleccionado para Ridge es: {lambda_optimo_ridge:.5f}")



# --- Generación de Box-plots --- 

# Los 'Cs' se ordenan de menor a mayor en el objeto CV, por lo que los lambdas correspondientes van de mayor a menor.
log_lambdas_sorted = np.log10(1 / lasso_cv_model.Cs_)

fig, axes = plt.subplots(1, 3, figsize=(24, 7), gridspec_kw={'width_ratios': [1, 1, 1]})

labels = [f'{l:.1f}' for l in log_lambdas_sorted]

# 1. Box-plot de error para LASSO (cada caja corresponde a un λ -> por eso transponemos)
axes[0].boxplot(lasso_errors.T.tolist())
axes[0].set_xticklabels(labels, rotation=45)
axes[0].set_title('Error de Clasificación (LASSO L1)', fontsize=14)
axes[0].set_xlabel('log10(λ)', fontsize=12)
axes[0].set_ylabel('Error de Clasificación (1 - Accuracy)', fontsize=12)
axes[0].grid(True, linestyle='--', alpha=0.6)

# 2. Box-plot de error para Ridge
axes[1].boxplot(ridge_errors.T.tolist())
axes[1].set_xticklabels(labels, rotation=45)
axes[1].set_title('Error de Clasificación (Ridge L2)', fontsize=14)
axes[1].set_xlabel('log10(λ)', fontsize=12)
axes[1].grid(True, linestyle='--', alpha=0.6)

# 3. (Opcional) Box-plot de proporción de ceros para LASSO
axes[2].boxplot(prop_zeros_lasso.T.tolist())
axes[2].set_xticklabels(labels, rotation=45)
axes[2].set_title('Proporción de Coeficientes Cero (LASSO L1)', fontsize=14)
axes[2].set_xlabel('log10(λ)', fontsize=12)
axes[2].set_ylabel('Proporción de Coeficientes = 0', fontsize=12)
axes[2].grid(True, linestyle='--', alpha=0.6)

# --- Añadir marcadores del λ óptimo en los boxplots ---
# posiciones en boxplot comienzan en 1
opt_log_lasso = np.log10(lambda_optimo_lasso)
opt_log_ridge = np.log10(lambda_optimo_ridge)

# encontrar la posición del tick más cercano
pos_lasso = int(np.argmin(np.abs(log_lambdas_sorted - opt_log_lasso))) + 1
pos_ridge = int(np.argmin(np.abs(log_lambdas_sorted - opt_log_ridge))) + 1

# líneas verticales y etiquetas
axes[0].axvline(pos_lasso, color='red', linestyle='--', linewidth=1.5)
axes[0].text(pos_lasso, axes[0].get_ylim()[1]*0.95, f'λ_opt={lambda_optimo_lasso:.2g}', color='red', ha='center', fontsize=10)

axes[1].axvline(pos_ridge, color='red', linestyle='--', linewidth=1.5)
axes[1].text(pos_ridge, axes[1].get_ylim()[1]*0.95, f'λ_opt={lambda_optimo_ridge:.2g}', color='red', ha='center', fontsize=10)


plt.suptitle('Distribución del Error y Selección de Variables vs. Penalidad λ', fontsize=16, y=1.02)
plt.tight_layout()
plt.show()





# --- Estimaciones comparativas: sin penalidad (statsmodels), L1 y L2 (sklearn) ---
# Usamos: X_train_const (con const) y X_train_dummies (sin const) y result (statsmodels ya ajustado)

# Valores de C seleccionados en CV (ya calculados)
C_l1 = lasso_cv_model.C_[0]
C_l2 = ridge_cv_model.C_[0]

# Preparar datos para penalizados: usar las dummies (sin const) y escalar
X_pen = X_train_dummies.copy()  # variables en la misma orden que el modelo sin penalidad (sin const)
y_pen = y_train  # target definido antes (train_2025_full[Y])

scaler_pen = StandardScaler()
X_pen_scaled = scaler_pen.fit_transform(X_pen)



# Ajuste L1 con C encontrado por CV
model_l1 = LogisticRegression(penalty='l1', C=C_l1, solver='liblinear', random_state=random_seed, max_iter=2000)
model_l1.fit(X_pen_scaled, y_pen)

# Ajuste L2 con C encontrado por CV
model_l2 = LogisticRegression(penalty='l2', C=C_l2, solver='liblinear', random_state=random_seed, max_iter=2000)
model_l2.fit(X_pen_scaled, y_pen)

# Desescalar coeficientes para volver a la escala original:
def unscale_coef(logreg, scaler, feature_names):
    b_scaled = logreg.coef_.ravel()
    mean = scaler.mean_
    scale = scaler.scale_
    # coef en escala original
    b_unscaled = b_scaled / scale
    # intercept en escala original
    intercept_unscaled = logreg.intercept_[0] - np.sum(b_scaled * mean / scale)
    s = pd.Series(index=list(feature_names) + ['const'], dtype=float)
    # poner coeficientes y const en el mismo orden que X_train_const (statsmodels)
    for i, fn in enumerate(feature_names):
        s[fn] = b_unscaled[i]
    s['const'] = intercept_unscaled
    return s

s_l1 = unscale_coef(model_l1, scaler_pen, X_pen.columns)
s_l2 = unscale_coef(model_l2, scaler_pen, X_pen.columns)

# Serie de coeficientes sin penalidad (statsmodels) ya disponible en `result.params`
s_unpen = result.params.copy()
# Asegurar mismo orden de índices: usar index de statsmodels (X_train_const.columns)
all_index = X_train_const.columns

# Construir DataFrame final con las tres columnas
df_coefs = pd.DataFrame(index=all_index)
df_coefs['Unpenalized'] = s_unpen.reindex(all_index)
df_coefs['L1 (λ^cv)'] = s_l1.reindex(all_index)
df_coefs['L2 (λ^cv)'] = s_l2.reindex(all_index)

# Formatear
df_coefs = df_coefs.round(5)

# Exportar a docx (tabla)
doc_coefs = Document()
doc_coefs.add_heading('Coeficientes: Sin penalidad vs L1 (λ^cv) vs L2 (λ^cv)', level=1)

t = doc_coefs.add_table(rows=1, cols=1 + df_coefs.shape[1])  # primera columna para nombre de variable
hdr = t.rows[0].cells
hdr[0].text = 'Variable'
for j, col in enumerate(df_coefs.columns, start=1):
    hdr[j].text = col

for var, row in df_coefs.iterrows():
    cells = t.add_row().cells
    cells[0].text = str(var)
    for j, val in enumerate(row, start=1):
        cells[j].text = f"{val:.5f}" if not pd.isna(val) else ""

# Ajustar tamaño de fuente
for row in t.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc_coefs.save('Coeficientes_Logit_Comparativo.docx')

# Guardar también como CSV por si se desea revisión rápida
df_coefs.to_csv('Coeficientes_Logit_Comparativo.csv')

# Mostrar resumen mínimo en consola
print("\nTabla de coeficientes creada: 'Coeficientes_Logit_Comparativo.docx' y .csv")
print(df_coefs.head(12))
